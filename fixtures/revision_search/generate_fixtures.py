"""Stage 7R.2: deterministic source-document fixtures for the
POLICY-RETENTION-001 revision-search benchmark.

Five REAL DOCX documents, one format only (this benchmark tests revision
authority, not format parity -- see Stage 6B/7A.1 for that). All five are
near-identical (same title, same section structure, same wording) except
the retention-period sentence itself, so that ordinary unfiltered vector
retrieval may legitimately rank an ineligible revision highly -- exactly
the condition Stage 7R.2 exists to prove authority-aware filtering
handles correctly.

No LLM calls, no network access, no randomness -- fixed timestamps, fixed
content, so re-running this script byte-for-byte reproduces the same five
files (and therefore the same source_document_sha256 for each).

Run: python fixtures/revision_search/generate_fixtures.py
"""

from __future__ import annotations

import json
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt

GENERATED_DIR = Path(__file__).resolve().parent / "generated"
FIXED_DOC_DATETIME = datetime(2026, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
FIXED_ZIP_DATETIME = (2026, 1, 1, 0, 0, 0)

# symbol -> (filename stem, retention-period sentence, extra status line or None)
REVISIONS: dict[str, tuple[str, str, str | None]] = {
    "v1": ("POLICY_RETENTION_v1", "3 years", None),
    "v2": ("POLICY_RETENTION_v2", "5 years", None),
    "v3": ("POLICY_RETENTION_v3", "7 years", None),
    "v4": ("POLICY_RETENTION_v4", "10 years", "Status: PROPOSED -- pending governance approval, not yet in effect."),
    "v5": ("POLICY_RETENTION_v5", "8 years", None),
}


def _normalize_zip_timestamps(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zf:
        entries = [(info, zf.read(info.filename)) for info in zf.infolist()]
    with zipfile.ZipFile(path, "w") as zf:
        for info, data in entries:
            new_info = zipfile.ZipInfo(info.filename, date_time=FIXED_ZIP_DATETIME)
            new_info.compress_type = info.compress_type
            new_info.external_attr = info.external_attr
            new_info.create_system = info.create_system
            zf.writestr(new_info, data)


def _set_core_properties(document: Document) -> None:
    props = document.core_properties
    props.author = "er-lab-ingestion-bench"
    props.last_modified_by = "er-lab-ingestion-bench"
    props.created = FIXED_DOC_DATETIME
    props.modified = FIXED_DOC_DATETIME
    props.title = ""
    props.subject = ""
    props.comments = ""
    props.category = ""
    props.revision = 1


def _build_docx(path: Path, retention_sentence: str, status_line: str | None) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    document.add_heading("Data Retention Policy — Customer Records", level=1)
    document.add_paragraph(
        "This policy governs the retention period for customer account records "
        "maintained by the organization."
    )
    document.add_heading("Retention Requirement", level=2)
    document.add_paragraph(
        f"Customer account records must be retained for a period of {retention_sentence} "
        "from the date of account closure, after which they must be securely destroyed in "
        "accordance with the organization's data disposal procedures."
    )
    if status_line:
        document.add_paragraph(status_line)
    document.add_heading("Scope", level=2)
    document.add_paragraph(
        "This policy applies to all customer-facing business units and any third-party "
        "data processors acting on the organization's behalf."
    )
    document.add_heading("Review", level=2)
    document.add_paragraph(
        "This policy is reviewed periodically by the governance committee to ensure "
        "continued alignment with applicable regulatory requirements."
    )

    _set_core_properties(document)
    document.save(str(path))
    _normalize_zip_timestamps(path)


def generate_all() -> dict[str, str]:
    """Returns {symbol: relative_path} for the five generated files."""
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    relative_paths: dict[str, str] = {}
    for symbol, (stem, retention_sentence, status_line) in REVISIONS.items():
        path = GENERATED_DIR / f"{stem}.docx"
        _build_docx(path, retention_sentence, status_line)
        relative_paths[symbol] = f"generated/{stem}.docx"
    return relative_paths


def main() -> None:
    relative_paths = generate_all()
    manifest = {"revisions": relative_paths}
    (Path(__file__).resolve().parent / "generation_manifest.json").write_text(
        json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    for symbol, rel in sorted(relative_paths.items()):
        print(f"{symbol}: {rel}")


if __name__ == "__main__":
    sys.exit(main())
