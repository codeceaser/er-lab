"""Stage 7B.0: deterministic source-document fixtures for the
cross-document relationship holdout benchmark.

An enterprise-like corpus in which the CURRENT relationship chain

    APP-224510 -> Payment Settlement -> Obligation O-31
               -> Control C-88 -> Procedure P-205 (current)

is deliberately DISTRIBUTED across five separate logical documents (one
hop each), so that no single chunk or document contains the assembled
multi-hop answer. Projection-neutral historical/draft/adjacent
distractors live in separate revisions or a separate document:

  - APP-224499  : retired application  (APP-PORTFOLIO revision 1)
  - Control C-88a: superseded control  (OBLIGATION-REGISTER rev1 + CONTROL-LIBRARY rev1)
  - Procedure P-204: historical procedure (CONTROL-LIBRARY rev1 + PROCEDURE-CATALOGUE rev1)
  - Control C-91 : draft/proposed control (CONTROL-LIBRARY rev3, never activated)
  - an unrelated but lexically similar "Payment Reconciliation" chain
    (ADJACENT-DOMAIN) using O-32 / C-77 / P-301 / APP-330012.

Only source facts naturally present in the documents are used. NO graph
nodes, edges, paths, expected answers, or benchmark labels appear in any
document -- authority (current/historical/draft) comes exclusively from
the Stage 7R registry/resolver at query time, and the relationship chain
is evaluation truth held in the contract, never in a chunk.

One format only (DOCX): this benchmark tests distributed relationships,
not format parity (see Stage 6B/7A.1 for that).

Deterministic: fixed timestamps, fixed content, no randomness -- re-running
byte-for-byte reproduces every file (and therefore every
source_document_sha256). The generated files are tracked in git; the
benchmark runner re-verifies each tracked file's bytes against
generation_manifest.json before use (fixtures.py).

Run: python fixtures/cross_document/generate_fixtures.py
"""

from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt

GENERATED_DIR = Path(__file__).resolve().parent / "generated"
FIXED_DOC_DATETIME = datetime(2026, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
FIXED_ZIP_DATETIME = (2026, 1, 1, 0, 0, 0)


# Each source file: symbol -> (filename stem, doc title, section heading, [body paragraphs]).
# One PARAGRAPH == one atomic relationship fact, so a single-fact document
# yields a single-fact chunk and no real-chain document ever pre-assembles
# a multi-hop answer. The ADJACENT-DOMAIN distractor deliberately DOES
# co-locate its whole (unrelated) chain, making it a dense lexical
# distractor without ever being a required answer.
SOURCE_FILES: dict[str, tuple[str, str, str, list[str]]] = {
    # APP-PORTFOLIO -- rev1 historical (retired app), rev2 current.
    "app_rev1": (
        "APP_PORTFOLIO_rev1", "Application Portfolio", "Registered Applications",
        ["Application APP-224499 supports the Payment Settlement business service."],
    ),
    "app_rev2": (
        "APP_PORTFOLIO_rev2", "Application Portfolio", "Registered Applications",
        ["Application APP-224510 supports the Payment Settlement business service."],
    ),
    # SERVICE-CATALOGUE -- single current revision.
    "svc_rev1": (
        "SERVICE_CATALOGUE_rev1", "Business Service Catalogue", "Governed Services",
        ["The Payment Settlement business service is governed by Obligation O-31."],
    ),
    # OBLIGATION-REGISTER -- rev1 historical (C-88a), rev2 current (C-88).
    "obl_rev1": (
        "OBLIGATION_REGISTER_rev1", "Obligation Register", "Obligation Coverage",
        ["Obligation O-31 is satisfied by Control C-88a."],
    ),
    "obl_rev2": (
        "OBLIGATION_REGISTER_rev2", "Obligation Register", "Obligation Coverage",
        ["Obligation O-31 is satisfied by Control C-88."],
    ),
    # CONTROL-LIBRARY -- rev1 historical (C-88a->P-204), rev2 current
    # (C-88->P-205), rev3 draft/proposed (C-91, never activated).
    "ctl_rev1": (
        "CONTROL_LIBRARY_rev1", "Control Library", "Control Implementations",
        ["Control C-88a is implemented through Procedure P-204."],
    ),
    "ctl_rev2": (
        "CONTROL_LIBRARY_rev2", "Control Library", "Control Implementations",
        ["Control C-88 is implemented through Procedure P-205."],
    ),
    "ctl_rev3": (
        "CONTROL_LIBRARY_rev3", "Control Library", "Control Implementations",
        ["Control C-91 is implemented through Procedure P-205."],
    ),
    # PROCEDURE-CATALOGUE -- rev1 historical (P-204 retired), rev2 current (P-205).
    "prc_rev1": (
        "PROCEDURE_CATALOGUE_rev1", "Procedure Catalogue", "Operating Procedures",
        ["Procedure P-204 is a retired operating procedure."],
    ),
    "prc_rev2": (
        "PROCEDURE_CATALOGUE_rev2", "Procedure Catalogue", "Operating Procedures",
        ["Procedure P-205 is the current operating procedure."],
    ),
    # ADJACENT-DOMAIN -- one unrelated but lexically similar chain,
    # deliberately co-located in a single document as a dense distractor.
    "adj_rev1": (
        "ADJACENT_DOMAIN_rev1", "Adjacent Domain Reference", "Payment Reconciliation Chain",
        [
            "Application APP-330012 supports the Payment Reconciliation business service.",
            "The Payment Reconciliation business service is governed by Obligation O-32.",
            "Obligation O-32 is satisfied by Control C-77.",
            "Control C-77 is implemented through Procedure P-301.",
            "Procedure P-301 is the current operating procedure for reconciliation.",
        ],
    ),
}


def _normalize_zip_timestamps(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zf:
        entries = [(info, zf.read(info.filename)) for info in zf.infolist()]
    with zipfile.ZipFile(path, "w") as zf:
        for info, data in entries:
            new_info = zipfile.ZipInfo(info.filename, date_time=FIXED_ZIP_DATETIME)
            new_info.compress_type = info.compress_type
            new_info.external_attr = info.external_attr
            new_info.create_system = info.create_system
            zf.writestr(new_info, data)


def _set_core_properties(document: Document) -> None:
    props = document.core_properties
    props.author = "er-lab-ingestion-bench"
    props.last_modified_by = "er-lab-ingestion-bench"
    props.created = FIXED_DOC_DATETIME
    props.modified = FIXED_DOC_DATETIME
    props.title = ""
    props.subject = ""
    props.comments = ""
    props.category = ""
    props.revision = 1


def _build_docx(path: Path, title: str, heading: str, paragraphs: list[str]) -> None:
    document = Document()
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading(title, level=1)
    document.add_heading(heading, level=2)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    _set_core_properties(document)
    document.save(str(path))
    _normalize_zip_timestamps(path)


def generate_all() -> dict[str, str]:
    """Returns {symbol: relative_path} for every generated file."""
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    relative_paths: dict[str, str] = {}
    for symbol, (stem, title, heading, paragraphs) in SOURCE_FILES.items():
        path = GENERATED_DIR / f"{stem}.docx"
        _build_docx(path, title, heading, paragraphs)
        relative_paths[symbol] = f"generated/{stem}.docx"
    return relative_paths


def main() -> None:
    relative_paths = generate_all()
    source_document_sha256 = {
        symbol: hashlib.sha256((Path(__file__).resolve().parent / rel).read_bytes()).hexdigest()
        for symbol, rel in relative_paths.items()
    }
    manifest = {"revisions": relative_paths, "source_document_sha256": source_document_sha256}
    (Path(__file__).resolve().parent / "generation_manifest.json").write_text(
        json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    for symbol in sorted(relative_paths):
        print(f"{symbol}: {relative_paths[symbol]} sha256={source_document_sha256[symbol]}")


if __name__ == "__main__":
    sys.exit(main())
